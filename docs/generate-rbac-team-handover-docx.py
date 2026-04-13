from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

root = Path(r"D:/OneDrive - linhdangdev/Documents/Obsidian Vault/ShucomAIOT")
out = root / "docs" / "shuncom-rbac-workflow-team-handover-detailed-v4.docx"


def set_font(style, name="Arial", size=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        style.font.size = Pt(size)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
    set_font(doc.styles[style_name])

set_font(doc.styles["Normal"], size=10.5)
set_font(doc.styles["Title"], size=20)
set_font(doc.styles["Heading 1"], size=15)
set_font(doc.styles["Heading 2"], size=12.5)
set_font(doc.styles["Heading 3"], size=11.5)

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = p.add_run("SHUNCOM RULR RBAC Workflow\nTài liệu handover cho Backend và Frontend")
r.bold = True
r.font.size = Pt(20)

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = p.add_run("Tài liệu chuẩn hóa role, scope, permission, workflow, API và UI để team triển khai đồng bộ.")
r.italic = True


doc.add_heading("3. Cấu trúc domain được quản lý", level=1)
doc.add_paragraph("Hierarchy chuẩn của hệ thống như sau:")
add_bullets(doc, [
    "Project → Area → Zone → Device",
    "Project là phạm vi quản trị gốc của Project Admin và Project Member.",
    "Area là cấp con trực tiếp dưới Project.",
    "Zone là cấp con trực tiếp dưới Area.",
    "Device là cấp con trực tiếp dưới Zone và là đối tượng vận hành chính.",
])
doc.add_paragraph("Các domain liên quan đi cùng workflow này gồm Rules, Alarms, Dashboard, Reports, Logs và Project Membership.")


doc.add_heading("4. Mô hình role", level=1)
add_table(
    doc,
    ["Role", "Phạm vi", "Quyền hạn", "Lưu ý triển khai"],
    [
        ["Manufacturer", "Toàn hệ thống", "Toàn quyền nền tảng", "Có thể vượt project boundary nhưng vẫn phải audit đầy đủ"],
        ["Project Admin", "Một hoặc nhiều project được giao", "Toàn quyền trong managed project", "Phải bị chặn khi đi ra ngoài project được giao"],
        ["Project Member", "Project hoặc sub-scope được cấp", "Delegated permissions", "Mọi action nhạy cảm phải check permission key + scope constraint"],
    ],
)


doc.add_heading("5. Nguyên tắc nghiệp vụ bắt buộc", level=1)
add_bullets(doc, [
    "Manufacturer là role duy nhất có full visibility và full permission trên toàn hệ thống.",
    "Project Admin không được thao tác ngoài phạm vi project được giao.",
    "Project Member không được tự nâng quyền.",
    "Project Member không được quản lý membership.",
    "Project Admin chỉ được cấp quyền cho member trong project mình quản lý.",
    "Tài khoản bị disable không được login.",
    "Frontend chỉ phản ánh quyền; backend là nơi chốt quyền thật.",
    "Mọi thay đổi permission, scope, membership phải có audit log.",
])


doc.add_heading("6. Mô hình ra quyết định quyền truy cập", level=1)
doc.add_paragraph("Backend cần thống nhất cùng một logic khi xử lý mọi action nhạy cảm.")
add_bullets(doc, [
    "Xác thực token hoặc session.",
    "Xác minh trạng thái tài khoản còn active.",
    "Resolve role hiện tại của user.",
    "Resolve phạm vi được cấp.",
    "Resolve permission key mà action hiện tại yêu cầu.",
    "Resolve đường dẫn resource: project → area → zone → device.",
    "Đối chiếu role + permission + scope.",
    "Nếu hợp lệ thì cho phép thực thi và ghi audit log.",
    "Nếu không hợp lệ thì trả mã lỗi phù hợp.",
])


doc.add_heading("7. Quy ước 401 và 403", level=1)
add_table(
    doc,
    ["HTTP code", "Khi dùng", "Ví dụ"],
    [
        ["401", "Chưa xác thực hoặc tài khoản không hợp lệ", "Token sai, token hết hạn, tài khoản bị disable"],
        ["403", "Đã xác thực nhưng không đủ quyền hoặc vượt scope", "Project Admin chạm project ngoài scope, Member gọi delete device khi chỉ có read"],
    ],
)


doc.add_heading("8. Mô hình scope", level=1)
add_table(
    doc,
    ["Loại scope", "Đối tượng", "Ý nghĩa", "Gợi ý triển khai"],
    [
        ["Global", "Manufacturer", "Nhìn thấy và thao tác toàn hệ thống", "Authorization layer + global navigation"],
        ["Managed project", "Project Admin", "Toàn quyền trong project được giao", "user_projects + project context"],
        ["Delegated project", "Project Member", "Bị giới hạn trong project được cấp", "project_member_permissions"],
        ["Area constrained", "Project Member", "Chỉ nhìn thấy area được cấp", "constraints JSON + filtered queries"],
        ["Zone constrained", "Project Member", "Chỉ nhìn thấy zone được cấp", "constraints JSON + filtered queries"],
        ["Device constrained", "Project Member", "Chỉ nhìn thấy device được cấp", "constraints JSON + filtered queries"],
    ],
)


doc.add_heading("9. Nhóm permission chính", level=1)
add_table(
    doc,
    ["Nhóm", "Manufacturer", "Project Admin", "Project Member", "Ghi chú"],
    [
        ["Project / Area / Zone", "Full", "Full trong scope", "Configurable", "FE tree phải đi đúng thứ tự Project → Area → Zone → Device"],
        ["Devices", "Full", "Full trong scope", "Configurable", "Bao gồm read, create, update, delete, execute, import, export"],
        ["Rules", "Full", "Full trong scope", "Configurable", "Bao gồm create, update, delete, enable, execute"],
        ["Alarms", "Full", "Full trong scope", "Configurable", "Bao gồm acknowledge và resolve"],
        ["Dashboard / Reports", "Full", "Full trong scope", "Configurable", "Có thể là read-only hoặc export-only"],
        ["Logs", "Full", "Full trong scope", "Configurable", "Có thể tách audit log và system log nếu cần"],
        ["Project membership", "Full", "Full trong scope", "No", "Project Member không chạm nhóm này"],
    ],
)


doc.add_heading("10. Permission matrix tham chiếu", level=1)
add_table(
    doc,
    ["Permission key", "Manufacturer", "Project Admin", "Project Member"],
    [
        ["dashboard.read", "Yes", "Yes", "Configurable"],
        ["dashboard.configure", "Yes", "Yes", "Configurable"],
        ["dashboard.export", "Yes", "Yes", "Configurable"],
        ["devices.read", "Yes", "Yes", "Configurable"],
        ["devices.create", "Yes", "Yes", "Configurable"],
        ["devices.update", "Yes", "Yes", "Configurable"],
        ["devices.delete", "Yes", "Yes", "Configurable"],
        ["devices.execute", "Yes", "Yes", "Configurable"],
        ["devices.import", "Yes", "Yes", "Configurable"],
        ["devices.export", "Yes", "Yes", "Configurable"],
        ["rules.read", "Yes", "Yes", "Configurable"],
        ["rules.create", "Yes", "Yes", "Configurable"],
        ["rules.update", "Yes", "Yes", "Configurable"],
        ["rules.delete", "Yes", "Yes", "Configurable"],
        ["rules.enable", "Yes", "Yes", "Configurable"],
        ["rules.execute", "Yes", "Yes", "Configurable"],
        ["alarms.read", "Yes", "Yes", "Configurable"],
        ["alarms.acknowledge", "Yes", "Yes", "Configurable"],
        ["alarms.resolve", "Yes", "Yes", "Configurable"],
        ["alarms.configure", "Yes", "Yes", "Configurable"],
        ["reports.read", "Yes", "Yes", "Configurable"],
        ["reports.generate", "Yes", "Yes", "Configurable"],
        ["reports.export", "Yes", "Yes", "Configurable"],
        ["logs.read", "Yes", "Yes", "Configurable"],
        ["logs.export", "Yes", "Yes", "Configurable"],
        ["members.read", "Yes", "Yes", "No"],
        ["members.create", "Yes", "Yes", "No"],
        ["members.update", "Yes", "Yes", "No"],
        ["members.delete", "Yes", "Yes", "No"],
    ],
)


doc.add_heading("11. Preset của Project Member", level=1)
doc.add_paragraph(
    "Viewer, Operator, Engineer, Analyst không phải top-level role mới. Đây là các preset quyền áp dụng cho Project Member."
)

doc.add_heading("11.1 Viewer", level=2)
add_bullets(doc, [
    "Mục đích: chỉ xem và theo dõi.",
    "Bao gồm: dashboard.read, devices.read, rules.read, reports.read, logs.read nếu được cho phép.",
    "Không bao gồm: execute, create, update, delete.",
    "Phù hợp cho: giám sát, quản lý xem số liệu, tài khoản chỉ quan sát.",
])

doc.add_heading("11.2 Operator", level=2)
add_bullets(doc, [
    "Mục đích: vận hành hằng ngày.",
    "Bao gồm: quyền Viewer + devices.execute + alarms.acknowledge + alarms.resolve.",
    "Không bao gồm mặc định: membership management, rule editing.",
    "Phù hợp cho: trực vận hành, điều khiển thiết bị, xử lý alarm.",
])

doc.add_heading("11.3 Engineer", level=2)
add_bullets(doc, [
    "Mục đích: cấu hình kỹ thuật và bảo trì automation.",
    "Bao gồm: quyền Operator + rules.create/update/enable + devices.update trong phạm vi được cấp + local-rule sync nếu cần.",
    "Không bao gồm mặc định: membership management.",
    "Phù hợp cho: kỹ sư triển khai, kỹ sư automation, kỹ sư bảo trì.",
])

doc.add_heading("11.4 Analyst", level=2)
add_bullets(doc, [
    "Mục đích: báo cáo và phân tích.",
    "Bao gồm: dashboard.read + reports.read/generate/export + logs.read/export nếu cần.",
    "Không bao gồm: execute command, sửa cấu hình, sửa rule.",
    "Phù hợp cho: phân tích KPI, báo cáo, audit nội bộ.",
])

doc.add_heading("11.5 Lưu ý triển khai", level=2)
add_bullets(doc, [
    "Role lưu trong hệ thống vẫn là ProjectMember.",
    "Preset chỉ là gói permission để cấp nhanh.",
    "Frontend có thể dùng preset label cho UX.",
    "Backend phải check permission key và scope thật, không check theo tên preset.",
])


doc.add_heading("12. Luồng đa project cho Project Admin", level=1)
add_bullets(doc, [
    "Sau khi login, hệ thống tải danh sách managed projects.",
    "Nếu chỉ có một project thì vào thẳng dashboard của project đó.",
    "Nếu có nhiều project thì hiển thị project selector.",
    "Frontend set current project context sau khi user chọn project.",
    "Backend vẫn phải tự kiểm tra target resource thuộc đúng project scope.",
    "Frontend phải xóa state cũ khi đổi project để tránh hiển thị sai dữ liệu.",
])


doc.add_heading("13. Luồng onboarding", level=1)
add_bullets(doc, [
    "Tạo user profile cơ bản.",
    "Gán role chuẩn.",
    "Nếu là Project Admin hoặc Project Member thì gán project scope.",
    "Nếu là Project Member thì chọn preset quyền.",
    "Nếu cần thì chỉnh thêm permission cụ thể.",
    "Nếu có hỗ trợ thì gán thêm constraint theo area, zone, device.",
    "Review trước khi tạo membership record.",
    "Gửi credential hoặc invitation.",
])


doc.add_heading("14. Phạm vi implement của Backend", level=1)
add_bullets(doc, [
    "PostgreSQL: users, roles, permissions, role_permissions, projects, user_projects, project_member_permissions, permission_templates.",
    "MongoDB #1: devices, device_telemetry, device_groups.",
    "MongoDB #2: audit_logs, system_events, alarm_history.",
    "Redis: permission cache, scope cache, queue support nếu cần.",
    "Module chính: auth, users, projects, membership, permissions, devices, rules, alarms, audit logs.",
    "Authorization middleware phải check role + permission key + scope trên mọi endpoint nhạy cảm.",
    "Cross-database validation phải lần ngược được từ device về zone, area, project.",
])


doc.add_heading("15. Phạm vi implement của Frontend", level=1)
add_bullets(doc, [
    "Màn chính: user list, create/edit user, membership detail, preset picker, permission preview, custom override panel, project selector.",
    "Menu, button, action phải hide hoặc disable theo permission.",
    "Tree và filter phải đi theo Project → Area → Zone → Device.",
    "Project Admin nhiều project phải có project selector rõ ràng.",
    "Project Member chỉ được thấy dữ liệu trong delegated scope.",
    "403 phải hiển thị thành lỗi quyền, không phải lỗi hệ thống chung chung.",
])


doc.add_heading("16. Phạm vi API", level=1)
add_bullets(doc, [
    "Authentication: /auth/login, /auth/logout, /auth/refresh.",
    "Users: /users CRUD.",
    "Project membership: list members, add member, update permissions, remove member.",
    "Devices: list, batch import, batch command, job status.",
    "Audit logs: list và export.",
])


doc.add_heading("17. Yêu cầu response sau login", level=1)
add_bullets(doc, [
    "user.id, user.username, user.display_name, user.role, user.status.",
    "Danh sách project user được truy cập.",
    "Current project nếu có auto-select.",
    "Permission summary để FE render UI.",
    "Constraint summary cho Project Member nếu có.",
])


doc.add_heading("18. Phạm vi audit log", level=1)
add_bullets(doc, [
    "Login success và failed login.",
    "Tạo, sửa, xóa user.",
    "Assign và revoke role.",
    "Add/remove member khỏi project.",
    "Đổi permission set.",
    "Device commands và batch actions nhạy cảm.",
    "Export, import, delete actions.",
    "Access denial nếu cần phục vụ điều tra.",
])


doc.add_heading("19. Use case cốt lõi", level=1)
doc.add_heading("19.1 Manufacturer truy cập toàn hệ thống", level=2)
add_bullets(doc, [
    "Actor: Manufacturer.",
    "Hành vi: mở dashboard tổng và đi vào mọi project.",
    "Kỳ vọng: thấy toàn bộ project, area, zone, device, rules, reports.",
])

doc.add_heading("19.2 Project Admin quản lý nhiều project", level=2)
add_bullets(doc, [
    "Actor: Project Admin.",
    "Tiền điều kiện: được gán Project A và Project B.",
    "Hành vi: switch sang Project B rồi update device trong Project B.",
    "Kỳ vọng: pass trong Project B, fail ở Project C nếu không được giao.",
])

doc.add_heading("19.3 Project Member kiểu Viewer", level=2)
add_bullets(doc, [
    "Actor: Project Member.",
    "Tiền điều kiện: dùng preset Viewer.",
    "Hành vi: xem dashboard, xem devices, thử execute command.",
    "Kỳ vọng: read pass, execute fail.",
])

doc.add_heading("19.4 Project Member kiểu Engineer trong scope giới hạn", level=2)
add_bullets(doc, [
    "Actor: Project Member.",
    "Tiền điều kiện: preset Engineer nhưng chỉ trong zone được gán.",
    "Hành vi: update rule trong zone được cấp và zone ngoài phạm vi.",
    "Kỳ vọng: trong scope pass, ngoài scope fail.",
])


doc.add_heading("20. Phạm vi test tối thiểu", level=1)
doc.add_heading("20.1 Backend", level=2)
add_bullets(doc, [
    "Manufacturer pass trên toàn bộ endpoint nhạy cảm.",
    "Project Admin chỉ pass trong managed project.",
    "Project Admin fail ngoài managed project.",
    "Project Member thiếu permission key thì nhận 403.",
    "Project Member có permission nhưng vượt scope thì nhận 403.",
    "Disabled user nhận 401 khi login.",
    "Permission change phải có hiệu lực ở request kế tiếp.",
    "Membership và permission update phải sinh audit log.",
])

doc.add_heading("20.2 Frontend", level=2)
add_bullets(doc, [
    "Viewer không thấy execute, edit, delete.",
    "Operator thấy execute nhưng không thấy rule edit mặc định.",
    "Project Admin nhiều project phải thấy project selector.",
    "Khi đổi project, state theo project phải được refresh đúng.",
    "403 phải hiển thị thành thông báo không đủ quyền.",
])

doc.add_heading("20.3 Integration", level=2)
add_bullets(doc, [
    "Frontend project context và backend scope validation phải nhất quán.",
    "Permission changes phải phản ánh sau khi refresh profile.",
    "Batch operations vẫn phải enforce scope đúng.",
    "Audit logs phải có actor, target, action, timestamp, project reference.",
])


doc.add_heading("21. Checklist chốt trước khi code", level=1)
add_bullets(doc, [
    "Chốt danh sách permission key.",
    "Chốt shape response của login/profile.",
    "Chốt constraints model cho area, zone, device.",
    "Chốt behavior cho Project Admin nhiều project.",
    "Chốt chiến lược current project state ở FE.",
    "Chốt danh sách action bắt buộc phải audit.",
    "Chốt semantics 401 và 403.",
])


doc.add_heading("22. Tóm tắt", level=1)
doc.add_paragraph(
    "Nguyên tắc quan trọng nhất là role không đủ để quyết định quyền. Mọi action nhạy cảm phải được đánh giá bằng role, permission key và scope. "
    "Project Member vẫn là một role duy nhất nhưng có nhiều preset quyền. Frontend phản ánh quyền; Backend cưỡng chế quyền."
)

doc.add_heading("23. Nguồn tài liệu", level=1)
doc.add_paragraph(
    "Tài liệu này được rút gọn và chuẩn hóa từ plans/reports/brainstormer-260413-1202-rbac-multi-project-workflow.md cùng các cập nhật hierarchy đã xác nhận."
)

doc.save(out)
print(out)
